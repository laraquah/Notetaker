import streamlit as st
import streamlit.components.v1 as components
import os
import shutil
import tempfile
from docx import Document
from docx.shared import Pt, Inches, RGBColor
import io
import time
import subprocess
import pickle
import json
import datetime
import pytz
import re
import requests
from requests_oauthlib import OAuth2Session

# --- FIX: ALLOW OAUTH TO RUN ON STREAMLIT CLOUD ---
os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'

# Import Google Cloud Libraries
from google.cloud import speech
from google.cloud import storage
import google.generativeai as genai

# Import Google Auth & Drive Libraries
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import Flow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload, MediaIoBaseDownload
from google.oauth2 import service_account

# -----------------------------------------------------
# 1. CONSTANTS & CONFIGURATION
# -----------------------------------------------------
st.set_page_config(layout="wide", page_title="AI Meeting Manager", page_icon="🤖")

# --- Load App Keys from Secrets ---
try:
    GCS_BUCKET_NAME = st.secrets["GCS_BUCKET_NAME"]
    GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
    GCP_SERVICE_ACCOUNT_JSON = json.loads(st.secrets["GCP_SERVICE_ACCOUNT_JSON"])
    GDRIVE_CLIENT_CONFIG = json.loads(st.secrets["GDRIVE_CLIENT_SECRET_JSON"])
    BASECAMP_CLIENT_ID = st.secrets["BASECAMP_CLIENT_ID"]
    BASECAMP_CLIENT_SECRET = st.secrets["BASECAMP_CLIENT_SECRET"]
    BASECAMP_ACCOUNT_ID = st.secrets["BASECAMP_ACCOUNT_ID"]
    
    # --- AUTO-LOGIN LOGIC ---
    STREAMLIT_APP_URL = st.secrets.get("STREAMLIT_APP_URL", None)
    
    if STREAMLIT_APP_URL:
        BASECAMP_REDIRECT_URI = STREAMLIT_APP_URL.rstrip("/")
        AUTO_LOGIN_MODE = True
    else:
        BASECAMP_REDIRECT_URI = "https://www.google.com"
        AUTO_LOGIN_MODE = False

except Exception as e:
    st.error(f"Secrets Configuration Error: {e}")
    st.stop()

# URLs
BASECAMP_AUTH_URL = "https://launchpad.37signals.com/authorization/new"
BASECAMP_TOKEN_URL = "https://launchpad.37signals.com/authorization/token"
BASECAMP_API_BASE = f"https://3.basecampapi.com/{BASECAMP_ACCOUNT_ID}"
BASECAMP_USER_AGENT = {"User-Agent": "AI Meeting Notes App (external-user)"}

# --- API CLIENTS SETUP ---
try:
    sa_creds = service_account.Credentials.from_service_account_info(GCP_SERVICE_ACCOUNT_JSON)
    storage_client = storage.Client(credentials=sa_creds)
    speech_client = speech.SpeechClient(credentials=sa_creds)
    
    genai.configure(api_key=GOOGLE_API_KEY)
    gemini_model = genai.GenerativeModel('gemini-2.5-flash-lite')
except Exception as e:
    st.error(f"System Error (AI Services): {e}")
    st.stop()

# =====================================================
# 2. HELPER FUNCTIONS
# =====================================================

def fetch_basecamp_name(token_dict):
    try:
        identity_url = "https://launchpad.37signals.com/authorization.json"
        headers = {"Authorization": f"Bearer {token_dict['access_token']}", "User-Agent": "AI Meeting Notes App"}
        response = requests.get(identity_url, headers=headers)
        if response.status_code == 200:
            data = response.json()
            return f"{data.get('identity', {}).get('first_name', '')} {data.get('identity', {}).get('last_name', '')}".strip()
    except: return ""
    return ""

def get_basecamp_session_user():
    if not st.session_state.basecamp_token: return None
    session = OAuth2Session(BASECAMP_CLIENT_ID, token=st.session_state.basecamp_token)
    session.headers.update(BASECAMP_USER_AGENT)
    return session

def upload_to_gcs(file_path, destination_blob_name):
    try:
        bucket = storage_client.bucket(GCS_BUCKET_NAME)
        blob = bucket.blob(destination_blob_name)
        blob.upload_from_filename(file_path, timeout=3600)
        return f"gs://{GCS_BUCKET_NAME}/{destination_blob_name}"
    except Exception as e:
        st.error(f"GCS Upload Error: {e}")
        return None

def get_or_create_folder(service, folder_name):
    try:
        query = f"mimeType='application/vnd.google-apps.folder' and name='{folder_name}' and trashed=false"
        results = service.files().list(q=query, fields="files(id)").execute()
        items = results.get('files', [])
        if items: return items[0]['id']
        else:
            file_metadata = {'name': folder_name, 'mimeType': 'application/vnd.google-apps.folder'}
            folder = service.files().create(body=file_metadata, fields='id').execute()
            return folder.get('id')
    except Exception as e: return None

def upload_to_drive_user(file_stream, file_name, target_folder_name):
    if not st.session_state.gdrive_creds: return None
    try:
        service = build("drive", "v3", credentials=st.session_state.gdrive_creds)
        folder_id = get_or_create_folder(service, target_folder_name)
        parents = [folder_id] if folder_id else []

        file_metadata = {"name": file_name, "parents": parents}
        media = MediaIoBaseUpload(file_stream, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        file = service.files().create(body=file_metadata, media_body=media, fields="id").execute()
        return file.get("id")
    except Exception as e:
        st.error(f"Google Drive Upload Error: {e}")
        return None

def save_analysis_data_to_drive(data_dict, filename):
    if not st.session_state.gdrive_creds: return None
    try:
        service = build("drive", "v3", credentials=st.session_state.gdrive_creds)
        folder_id = get_or_create_folder(service, "Meeting_Data")
        if not folder_id: return None

        json_str = json.dumps(data_dict, indent=2)
        fh = io.BytesIO(json_str.encode('utf-8'))
        file_metadata = {"name": filename, "parents": [folder_id]}
        media = MediaIoBaseUpload(fh, mimetype='application/json')
        service.files().create(body=file_metadata, media_body=media, fields="id").execute()
        return True
    except: return False

def list_past_meetings():
    if not st.session_state.gdrive_creds: return []
    try:
        service = build("drive", "v3", credentials=st.session_state.gdrive_creds)
        folder_id = get_or_create_folder(service, "Meeting_Data")
        if not folder_id: return []
        query = f"'{folder_id}' in parents and mimeType='application/json' and trashed=false"
        results = service.files().list(q=query, fields="files(id, name, createdTime)", orderBy="createdTime desc").execute()
        return results.get('files', [])
    except: return []

def load_meeting_data(file_id):
    if not st.session_state.gdrive_creds: return None
    try:
        service = build("drive", "v3", credentials=st.session_state.gdrive_creds)
        request = service.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while done is False:
            status, done = downloader.next_chunk()
        fh.seek(0)
        return json.load(fh)
    except: return None

# --- Basecamp Helpers ---
def get_basecamp_projects(_session):
    try:
        response = _session.get(f"{BASECAMP_API_BASE}/projects.json")
        response.raise_for_status()
        return sorted([(p['name'], p['id']) for p in response.json() if p['status'] == 'active'], key=lambda x: x[0])
    except: return []

def get_project_tools(_session, project_id):
    try:
        response = _session.get(f"{BASECAMP_API_BASE}/projects/{project_id}.json")
        response.raise_for_status()
        return response.json().get('dock', [])
    except: return []

def get_todolists(_session, todoset_id, project_id):
    try:
        url = f"{BASECAMP_API_BASE}/buckets/{project_id}/todosets/{todoset_id}/todolists.json"
        resp = _session.get(url)
        return sorted([(t['title'], t['id']) for t in resp.json()], key=lambda x: x[0])
    except: return []

def upload_bc_attachment(_session, file_bytes, file_name):
    try:
        headers = _session.headers.copy()
        headers.update({'Content-Type': 'application/octet-stream', 'Content-Length': str(len(file_bytes))})
        resp = _session.post(f"{BASECAMP_API_BASE}/attachments.json?name={file_name}", data=file_bytes, headers=headers)
        return resp.json()['attachable_sgid']
    except Exception as e:
        st.error(f"Basecamp Upload Error: {e}")
        return None

def post_to_basecamp(_session, project_id, tool_type, tool_id, sub_id, title, content, attachment_sgid):
    try:
        attach_html = f'<bc-attachment sgid="{attachment_sgid}"></bc-attachment>' if attachment_sgid else ""
        
        if tool_type == "To-dos":
            url = f"{BASECAMP_API_BASE}/buckets/{project_id}/todolists/{sub_id}/todos.json"
            payload = {"content": title, "description": content + attach_html}
        elif tool_type == "Message Board":
            url = f"{BASECAMP_API_BASE}/buckets/{project_id}/message_boards/{tool_id}/messages.json"
            payload = {"subject": title, "content": content + attach_html, "status": "active"}
        elif tool_type == "Docs & Files":
            url = f"{BASECAMP_API_BASE}/buckets/{project_id}/vaults/{tool_id}/uploads.json"
            payload = {"attachable_sgid": attachment_sgid, "base_name": title, "content": content}

        resp = _session.post(url, json=payload)
        resp.raise_for_status()
        return True
    except Exception as e:
        st.error(f"Basecamp Post Error: {e}")
        return False

# --- AI Analysis ---
def get_visual_metadata(file_path):
    if shutil.which("ffmpeg") is None: return None
    thumbnail_path = "temp_thumb.jpg"
    result_data = {"datetime_sg": None, "duration": 0, "title": "Meeting_Minutes", "venue": ""}
    try:
        # Duration
        cmd = ["ffprobe", "-v", "error", "-show_entries", "format=duration", "-of", "default=noprint_wrappers=1:nokey=1", file_path]
        res = subprocess.run(cmd, capture_output=True, text=True)
        if res.returncode == 0: result_data["duration"] = float(res.stdout.strip())

        # Vision
        subprocess.run(['ffmpeg', '-i', file_path, '-ss', '00:00:01', '-vframes', '1', '-q:v', '2', '-y', thumbnail_path], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

        if os.path.exists(thumbnail_path):
            vision = genai.GenerativeModel('gemini-2.5-flash-lite')
            with open(thumbnail_path, "rb") as img: img_data = img.read()
            prompt = """Analyze this meeting screenshot. Return JSON: { "datetime": "YYYY-MM-DD HH:MM", "title": "Center Text", "venue": "Corner Text" }. If not found, use "None"."""
            resp = vision.generate_content([{'mime_type': 'image/jpeg', 'data': img_data}, prompt])
            try:
                data = json.loads(resp.text.strip().replace("```json", "").replace("```", ""))
                if data.get("title") != "None": result_data["title"] = data["title"].replace(" ", "_")
                if data.get("venue") != "None": result_data["venue"] = data["venue"]
                if data.get("datetime") != "None":
                    dt = datetime.datetime.strptime(data["datetime"], "%Y-%m-%d %H:%M").replace(tzinfo=datetime.timezone.utc)
                    result_data["datetime_sg"] = dt.astimezone(pytz.timezone('Asia/Singapore'))
            except: pass
    except: pass
    finally:
        if os.path.exists(thumbnail_path): os.remove(thumbnail_path)
    return result_data

def get_structured_notes_google(audio_file_path, file_name, participants_context):
    try:
        with st.spinner(f"Converting {file_name}..."):
            base_name = os.path.splitext(audio_file_path)[0]
            flac_file_path = f"{base_name}.flac"
            subprocess.run(["ffmpeg", "-i", audio_file_path, "-vn", "-acodec", "flac", "-y", flac_file_path], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

        with st.spinner(f"Uploading to Google Cloud..."):
            gcs_uri = upload_to_gcs(flac_file_path, f"{os.path.splitext(file_name)[0]}.flac")
            if not gcs_uri: return {"error": "Upload failed."}

        progress_bar = st.progress(0, text="Transcribing...")
        audio = speech.RecognitionAudio(uri=gcs_uri)
        config = speech.RecognitionConfig(encoding=speech.RecognitionConfig.AudioEncoding.FLAC, language_code="en-US", enable_automatic_punctuation=True, use_enhanced=True, model="video", diarization_config=speech.SpeakerDiarizationConfig(enable_speaker_diarization=True, min_speaker_count=2, max_speaker_count=6))
        operation = speech_client.long_running_recognize(config=config, audio=audio)
        
        while not operation.done():
            if operation.metadata: progress_bar.progress(operation.metadata.progress_percent, text=f"Transcribing: {operation.metadata.progress_percent}%")
            time.sleep(2)

        progress_bar.progress(100, text="Done!")
        response = operation.result(timeout=3600)
        progress_bar.empty()

        if not response.results: return {"error": "Transcription failed."}

        result = response.results[-1]
        words = result.alternatives[0].words
        full_transcript = ""
        curr = -1
        for w in words:
            if w.speaker_tag != curr:
                curr = w.speaker_tag
                full_transcript += f"\n\nSpeaker {curr}: "
            full_transcript += w.word + " "

        with st.spinner("Analyzing with Gemini..."):
            # --- ROBUST PROMPT TO FORCE CONTENT INTO FIELDS ---
            prompt = f"""
            You are an expert meeting secretary. Context: {participants_context}
            Transcript: {full_transcript}
            
            TASKS:
            1. Identify speakers using context.
            2. Extract Sections using these EXACT headers:
            
            ## OVERVIEW ##
            [Brief summary of WHO met and WHAT was discussed (2-3 sentences).]
            
            ## DISCUSSION ##
            [Detailed bullet points with headers. Be comprehensive.]
            
            ## NEXT STEPS ##
            List specific actionable items. **IMPORTANT: Merge any Client Requests into this list as actions for the appropriate person.**
            FORMAT:
            * **Action:** [Specific Task] (Assigned to: [Name]) - Deadline: [Time if mentioned]
            """
            text = gemini_model.generate_content(prompt).text
            
            # --- ROBUST REGEX PARSER (FIXES EMPTY FIELDS) ---
            overview = ""
            discussion = ""
            next_steps = ""
            
            try:
                # Use Regex to capture content between headers, ignoring casing/spacing
                ov_match = re.search(r'##\s*OVERVIEW\s*##(.*?)(?=##\s*DISCUSSION|##\s*NEXT STEPS|$)', text, re.DOTALL | re.IGNORECASE)
                disc_match = re.search(r'##\s*DISCUSSION\s*##(.*?)(?=##\s*NEXT STEPS|$)', text, re.DOTALL | re.IGNORECASE)
                ns_match = re.search(r'##\s*NEXT STEPS\s*##(.*)', text, re.DOTALL | re.IGNORECASE)
                
                if ov_match: overview = ov_match.group(1).strip()
                if disc_match: discussion = disc_match.group(1).strip()
                if ns_match: next_steps = ns_match.group(1).strip()

                # Safety Fallback
                if not overview and not discussion: discussion = text
            except: 
                discussion = text

            return {"overview": overview, "discussion": discussion, "next_steps": next_steps, "full_transcript": full_transcript}
    except Exception as e: return {"error": str(e)}
    finally:
        try: bucket = storage_client.bucket(GCS_BUCKET_NAME); bucket.blob(f"{os.path.splitext(file_name)[0]}.flac").delete()
        except: pass

# --- Markdown Parsers ---
def _add_rich_text(paragraph, text):
    """Parses markdown-style **bold** text and applies Word formatting."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

def safe_apply_style(paragraph, style_name, fallback_prefix=""):
    try: paragraph.style = style_name
    except KeyError: 
        if fallback_prefix: paragraph.text = fallback_prefix + paragraph.text

def add_formatted_text(cell, text):
    """Enhanced parser to handle bullets and bold text correctly in Word table cells."""
    cell.text = ""
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line: continue
        
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        
        if line.startswith('##'):
            clean_title = line.lstrip('#').strip()
            run = p.add_run(clean_title)
            run.underline = True
            run.bold = True
            p.paragraph_format.space_before = Pt(8)
        elif line.startswith('*') or line.startswith('-'):
            clean_text = line.lstrip('*- ').strip()
            safe_apply_style(p, 'List Bullet', "• ") # SAFE STYLE APPLY
            _add_rich_text(p, clean_text)
            p.paragraph_format.left_indent = Inches(0.15)
        else:
            _add_rich_text(p, line)

def add_markdown_to_doc(doc, text):
    lines = text.split('\n')
    table_row = re.compile(r"^\|(.+)\|")
    table_sep = re.compile(r"^\|[-:| ]+\|")
    table_data = []
    in_table = False

    for line in lines:
        stripped = line.strip()
        if table_row.match(stripped):
            if not table_sep.match(stripped):
                table_data.append([c.strip() for c in stripped.strip('|').split('|')])
            in_table = True
            continue
        elif in_table:
            if table_data:
                t = doc.add_table(rows=len(table_data), cols=max(len(r) for r in table_data))
                t.style = 'Table Grid'
                for i, row in enumerate(table_data):
                    for j, val in enumerate(row):
                        if j < len(t.rows[i].cells):
                            p = t.rows[i].cells[j].paragraphs[0]
                            _add_rich_text(p, val)
                            if i == 0: 
                                for run in p.runs: run.bold = True
                doc.add_paragraph("")
                table_data = []
            in_table = False

        if stripped.startswith('##'):
            doc.add_heading(stripped.lstrip('#').strip(), level=2)
        elif stripped.startswith('*') or stripped.startswith('-'):
            p = doc.add_paragraph(style='List Bullet')
            _add_rich_text(p, stripped.lstrip('*- ').strip())
        elif stripped:
            p = doc.add_paragraph()
            _add_rich_text(p, stripped)
    
    if table_data: 
        t = doc.add_table(rows=len(table_data), cols=max(len(r) for r in table_data))
        t.style = 'Table Grid'
        for i, row in enumerate(table_data):
            for j, val in enumerate(row):
                if j < len(t.rows[i].cells):
                    p = t.rows[i].cells[j].paragraphs[0]
                    _add_rich_text(p, val)
                    if i == 0: 
                        for run in p.runs: run.bold = True

# -----------------------------------------------------
# 3. STATE & LOGIN
# -----------------------------------------------------
if 'gdrive_creds' not in st.session_state: st.session_state.gdrive_creds = None
if 'basecamp_token' not in st.session_state: st.session_state.basecamp_token = None
if 'user_real_name' not in st.session_state: st.session_state.user_real_name = ""
if 'detected_date' not in st.session_state: st.session_state.detected_date = None
if 'detected_time' not in st.session_state: st.session_state.detected_time = None
if 'detected_title' not in st.session_state: st.session_state.detected_title = "Meeting_Minutes"
if 'detected_venue' not in st.session_state: st.session_state.detected_venue = ""
if 'auto_client_reps' not in st.session_state: st.session_state.auto_client_reps = ""
if 'auto_ifoundries_reps' not in st.session_state: st.session_state.auto_ifoundries_reps = ""

# Re-hydrate Google
if 'gdrive_creds_json' in st.session_state and not st.session_state.gdrive_creds:
    try: st.session_state.gdrive_creds = Credentials.from_authorized_user_info(json.loads(st.session_state.gdrive_creds_json))
    except: st.session_state.gdrive_creds_json = None

# Basecamp Auto-Login
if AUTO_LOGIN_MODE and "code" in st.query_params and not st.session_state.basecamp_token:
    auth_code = st.query_params["code"]
    try:
        payload = {
            "type": "web_server",
            "client_id": BASECAMP_CLIENT_ID,
            "client_secret": BASECAMP_CLIENT_SECRET,
            "redirect_uri": BASECAMP_REDIRECT_URI,
            "code": auth_code
        }
        response = requests.post(BASECAMP_TOKEN_URL, data=payload)
        if response.status_code != 200:
            st.error(f"Basecamp Login Error ({response.status_code}): {response.text}")
        else:
            token = response.json()
            st.session_state.basecamp_token = token
            real_name = fetch_basecamp_name(token)
            if real_name: st.session_state.user_real_name = real_name
            st.toast("✅ Basecamp Login Successful!", icon="🎉")
            st.query_params.clear()
            time.sleep(1)
            st.rerun()
    except Exception as e:
        st.error(f"Auto-login system error: {e}")

# -----------------------------------------------------
# 4. SIDEBAR
# -----------------------------------------------------
with st.sidebar:
    st.title("🔐 Login")
    st.markdown("### Step 1: Basecamp")
    if st.session_state.basecamp_token:
        st.success(f"✅ Connected: {st.session_state.user_real_name}")
        if st.button("Logout Basecamp"):
            st.session_state.basecamp_token = None; st.session_state.user_real_name = ""; st.rerun()
    else:
        bc = OAuth2Session(BASECAMP_CLIENT_ID, redirect_uri=BASECAMP_REDIRECT_URI)
        url, _ = bc.authorization_url(BASECAMP_AUTH_URL, type="web_server")
        if AUTO_LOGIN_MODE: st.link_button("Login to Basecamp", url, type="primary")
        else: 
            st.markdown(f"[Authorize]({url})"); c = st.text_input("Code")
            if c: 
                st.session_state.basecamp_token = requests.post(BASECAMP_TOKEN_URL, data={"type":"web_server","client_id":BASECAMP_CLIENT_ID,"client_secret":BASECAMP_CLIENT_SECRET,"redirect_uri":BASECAMP_REDIRECT_URI,"code":c}).json()
                st.session_state.user_real_name = fetch_basecamp_name(st.session_state.basecamp_token)
                st.rerun()

    st.divider()
    st.markdown("### Step 2: Google Drive")
    if not st.session_state.basecamp_token: st.info("Please login to Basecamp first.")
    else:
        if st.session_state.gdrive_creds:
            st.success("✅ Connected")
            if st.button("Logout Drive"):
                st.session_state.gdrive_creds = None; st.session_state.gdrive_creds_json = None; st.rerun()
        else:
            f = Flow.from_client_config(GDRIVE_CLIENT_CONFIG, scopes=["https://www.googleapis.com/auth/drive"], redirect_uri="urn:ietf:wg:oauth:2.0:oob")
            url, _ = f.authorization_url(prompt='consent')
            st.link_button("Login to Drive", url)
            c = st.text_input("Paste Drive Code")
            if c:
                f.fetch_token(code=c)
                st.session_state.gdrive_creds = f.credentials
                st.session_state.gdrive_creds_json = f.credentials.to_json()
                st.rerun()

if not (st.session_state.basecamp_token and st.session_state.gdrive_creds):
    st.title("🔒 Access Restricted"); st.warning("Please login to both services."); st.stop()

# -----------------------------------------------------
# 8. MAIN UI
# -----------------------------------------------------
if 'ai_results' not in st.session_state: st.session_state.ai_results = {}
if 'chat_history' not in st.session_state: st.session_state.chat_history = []
if 'saved_participants_input' not in st.session_state: st.session_state.saved_participants_input = ""

st.title("🤖 AI Meeting Manager")

tab1, tab2, tab3, tab4 = st.tabs(["1. Analyze", "2. Review & Export", "3. Chat", "4. History"])

with tab1:
    st.header("1. Analyze Audio")
    participants = st.text_area("Participants", "Client (Client)\niFoundries (iFoundries)")
    up = st.file_uploader("Upload", type=['mp3','mp4','m4a','wav'])
    
    if st.button("Analyze"):
        if up:
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{up.name.split('.')[-1]}") as tmp:
                tmp.write(up.getvalue()); path = tmp.name
            
            st.session_state.chat_history = []
            st.session_state.saved_participants_input = participants
            
            # Parse Reps
            cl = [l.replace("(Client)","").strip() for l in participants.split('\n') if "(Client)" in l]
            il = [l.replace("(iFoundries)","").strip() for l in participants.split('\n') if "(iFoundries)" in l]
            st.session_state.auto_client_reps = "\n".join(cl)
            st.session_state.auto_ifoundries_reps = ", ".join(il)

            # Metadata
            with st.spinner("Extracting Metadata..."):
                meta = get_visual_metadata(path)
                if meta['datetime_sg']:
                    st.session_state.detected_date = meta['datetime_sg'].date()
                    end = meta['datetime_sg'] + datetime.timedelta(seconds=meta['duration'])
                    st.session_state.detected_time = f"{meta['datetime_sg'].strftime('%I:%M %p')} - {end.strftime('%I:%M %p')}"
                    if meta['title']: st.session_state.detected_title = meta['title']
                    if meta['venue']: st.session_state.detected_venue = meta['venue']
            
            # Analyze
            res = get_structured_notes_google(path, up.name, participants)
            if "error" in res: st.error(res["error"])
            else:
                st.session_state.ai_results = res
                # Auto-Save
                ts = datetime.datetime.now().strftime("%Y%m%d_%H%M")
                save_data = {
                    "ai_results": res, 
                    "participants": participants, 
                    "date": str(datetime.datetime.now()),
                    "chat_history": [],
                    "detected_title": st.session_state.detected_title
                }
                save_analysis_data_to_drive(save_data, f"Data_{up.name}_{ts}.json")
                st.success("Done! Check Review tab.")

with tab2:
    st.header("2. Review")
    d_date = st.session_state.detected_date if st.session_state.detected_date else datetime.date.today()
    d_time = st.session_state.detected_time if st.session_state.detected_time else datetime.datetime.now().strftime("%I:%M %p")
    
    c1, c2 = st.columns(2)
    date = c1.date_input("Date", d_date)
    venue = c1.text_input("Venue", st.session_state.detected_venue)
    crep = c1.text_input("Client Reps", st.session_state.auto_client_reps)
    absent = c1.text_input("Absent")
    
    time_str = c2.text_input("Time", d_time)
    prep = c2.text_input("Prepared By", st.session_state.user_real_name)
    irep = c2.text_input("iFoundries Reps", st.session_state.auto_ifoundries_reps)
    
    st.subheader("Content")
    overview = st.text_area("Overview (Green Box)", st.session_state.ai_results.get("overview", ""))
    disc = st.text_area("Discussion", st.session_state.ai_results.get("discussion", ""), height=300)
    next_s = st.text_area("Next Steps (Includes Client Requests)", st.session_state.ai_results.get("next_steps", ""), height=200)
    
    st.divider()
    do_d = st.checkbox("Upload to Drive", True)
    do_b = st.checkbox("Upload to Basecamp", False)
    
    pid, tool, tid, subid, btitle, bdesc = None, None, None, None, "", ""
    
    if do_b:
        sess = get_basecamp_session_user()
        projs = get_basecamp_projects(sess)
        pname = st.selectbox("Project", [p[0] for p in projs])
        if pname:
            pid = next(p[1] for p in projs if p[0]==pname)
            tool = st.selectbox("Where to post?", ["To-dos", "Message Board", "Docs & Files"])
            dock = get_project_tools(sess, pid)
            
            if tool == "To-dos":
                tid = next((t['id'] for t in dock if t['name']=='todoset'), None)
                lists = get_todolists(sess, tid, pid)
                lname = st.selectbox("List", [l[0] for l in lists])
                if lname: subid = next(l[1] for l in lists if l[0]==lname); btitle = st.text_input("Title", f"Minutes - {date}")
            elif tool == "Message Board":
                tid = next((t['id'] for t in dock if t['name']=='message_board'), None)
                btitle = st.text_input("Subject", f"Minutes - {date}")
            elif tool == "Docs":
                tid = next((t['id'] for t in dock if t['name']=='vault'), None)
                btitle = st.text_input("File Name", f"Minutes_{date}.docx")

    if st.button("Generate"):
        doc = Document("Minutes Of Meeting - Template.docx")
        
        # 1. Title Replace
        for p in doc.paragraphs: 
            if "[Title]" in p.text: p.text = p.text.replace("[Title]", st.session_state.detected_title.replace("_"," "))
        
        t = doc.tables[0]
        t.cell(1,1).text = str(date)
        t.cell(2,1).text = str(time_str)
        t.cell(3,1).text = venue
        
        # Force correct format: Name (Client) / Name (iFoundries)
        crep_final = crep if "(Client)" in crep else f"{crep} (Client)"
        irep_final = irep if "(iFoundries)" in irep else f"{irep} (iFoundries)"
        
        t.cell(4,1).text = crep_final
        t.cell(4,2).text = irep_final
        t.cell(5,1).text = absent
        
        # 2. Content
        t2 = doc.tables[1]
        t2.cell(1,1).text = overview # Green Box
        add_formatted_text(t2.cell(2,1), disc)
        add_formatted_text(t2.cell(4,1), next_s)
        
        # 3. Adjourned
        try: end_t = time_str.split('-')[1].strip()
        except: end_t = "Unknown"
        t2.cell(5,1).text = f"Meeting adjourned at {end_t}"
        
        doc.paragraphs[-1].text = f"Prepared by: {prep}"
        
        b = io.BytesIO(); doc.save(b); b.seek(0)
        fn = f"{st.session_state.detected_title}_{date}.docx"
        
        if do_d: upload_to_drive_user(b, fn, "Meeting Notes")
        if do_b and pid:
            sgid = upload_bc_attachment(sess, b.getvalue(), fn)
            post_to_basecamp(sess, pid, tool, tid, subid, btitle, "Attached.", sgid)
        
        st.success("Done!")
        st.download_button("Download", b, fn)

with tab3:
    st.header("💬 Chat")
    if st.button("💾 Save Chat to Drive"):
        d = Document()
        d.add_heading(f"Chat Log - {datetime.date.today()}", 0)
        for m in st.session_state.chat_history:
            p = d.add_paragraph()
            role = "AI" if m['role']=='assistant' else "User"
            p.add_run(f"{role}: ").bold = True
            add_markdown_to_doc(d, m['content'])
            d.add_paragraph("_"*30)
        b = io.BytesIO(); d.save(b); b.seek(0)
        upload_to_drive_user(b, f"Chat_{st.session_state.detected_title}.docx", "Chats")
        st.success("Saved!")

    # Chat Box
    box = st.container(height=500)
    with box:
        for m in st.session_state.chat_history:
            st.chat_message(m["role"], avatar="👤" if m["role"]=="user" else "🤖").markdown(m["content"])
    
    if p := st.chat_input("Ask a question..."):
        st.session_state.chat_history.append({"role":"user", "content":p})
        box.chat_message("user", avatar="👤").markdown(p)
        
        with box.chat_message("assistant", avatar="🤖"):
            with st.spinner("Thinking..."):
                prompt = f"""
                Role: Secretary.
                Context: {st.session_state.saved_participants_input}
                Transcript: {st.session_state.ai_results.get('full_transcript','')}
                Question: {p}
                Rules: Professional voice. Use real names. Concise.
                """
                resp = gemini_model.generate_content(prompt).text
                st.markdown(resp)
                st.session_state.chat_history.append({"role":"assistant", "content":resp})

with tab4:
    st.header("📂 History")
    if st.button("Refresh"): st.rerun()
    files = list_past_meetings()
    if files:
        sel = st.selectbox("Select Meeting", [f['name'] for f in files])
        if st.button("Load"):
            fid = next(f['id'] for f in files if f['name'] == sel)
            d = load_meeting_data(fid)
            if d:
                st.session_state.ai_results = d.get("ai_results", {})
                st.session_state.saved_participants_input = d.get("participants", "")
                # Restore Chat History!
                st.session_state.chat_history = d.get("chat_history", [])
                st.session_state.detected_title = d.get("detected_title", "Meeting")
                
                # Restore Reps
                p_input = st.session_state.saved_participants_input
                c_list = [l.replace("(Client)","").strip() for l in p_input.split('\n') if "(Client)" in l]
                i_list = [l.replace("(iFoundries)","").strip() for l in p_input.split('\n') if "(iFoundries)" in l]
                st.session_state.auto_client_reps = "\n".join(c_list)
                st.session_state.auto_ifoundries_reps = ", ".join(i_list)

                st.success("Loaded! Check Tab 2 and 3.")
                time.sleep(1); st.rerun()
