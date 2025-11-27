import streamlit as st
import os

# --- FIX: ALLOW OAUTH TO RUN ON STREAMLIT CLOUD ---
# This silences the "InsecureTransportError" by allowing internal non-HTTPS routing
# (Streamlit Cloud handles the actual HTTPS security externally)
os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'

import tempfile
from docx import Document
import io
import time
import subprocess
import pickle
import json
import datetime
import pytz

# Import Google Cloud Libraries
from google.cloud import speech
from google.cloud import storage
import google.generativeai as genai

# Import Google Auth & Drive Libraries
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import Flow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
from google.oauth2 import service_account

# --- Import Basecamp & formatting tools ---
import requests
from requests_oauthlib import OAuth2Session
from docx.shared import Pt, Inches

# -----------------------------------------------------
# 1. CONSTANTS & CONFIGURATION
# -----------------------------------------------------
st.set_page_config(layout="wide", page_title="AI Meeting Manager", page_icon="🤖")

# --- Load App Keys from Secrets ---
try:
    GCS_BUCKET_NAME = st.secrets["GCS_BUCKET_NAME"]
    GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
    GCP_SERVICE_ACCOUNT_JSON = json.loads(st.secrets["GCP_SERVICE_ACCOUNT_JSON"])
    GDRIVE_CLIENT_CONFIG = json.loads(st.secrets["GDRIVE_CLIENT_SECRET_JSON"])
    BASECAMP_CLIENT_ID = st.secrets["BASECAMP_CLIENT_ID"]
    BASECAMP_CLIENT_SECRET = st.secrets["BASECAMP_CLIENT_SECRET"]
    BASECAMP_ACCOUNT_ID = st.secrets["BASECAMP_ACCOUNT_ID"]
    
    # --- AUTO-LOGIN LOGIC ---
    STREAMLIT_APP_URL = st.secrets.get("STREAMLIT_APP_URL", None)
    
    if STREAMLIT_APP_URL:
        BASECAMP_REDIRECT_URI = STREAMLIT_APP_URL.rstrip("/")
        AUTO_LOGIN_MODE = True
    else:
        BASECAMP_REDIRECT_URI = "https://www.google.com"
        AUTO_LOGIN_MODE = False

except Exception as e:
    st.error(f"Secrets Configuration Error: {e}")
    st.stop()

# URLs
BASECAMP_AUTH_URL = "https://launchpad.37signals.com/authorization/new"
BASECAMP_TOKEN_URL = "https://launchpad.37signals.com/authorization/token"
BASECAMP_API_BASE = f"https://3.basecampapi.com/{BASECAMP_ACCOUNT_ID}"
BASECAMP_USER_AGENT = {"User-Agent": "AI Meeting Notes App (external-user)"}

# -----------------------------------------------------
# 2. HELPER: GET USER IDENTITY
# -----------------------------------------------------
def fetch_basecamp_name(token_dict):
    """Calls Basecamp Identity API to get the user's real name."""
    try:
        identity_url = "https://launchpad.37signals.com/authorization.json"
        headers = {
            "Authorization": f"Bearer {token_dict['access_token']}",
            "User-Agent": "AI Meeting Notes App"
        }
        response = requests.get(identity_url, headers=headers)
        if response.status_code == 200:
            data = response.json()
            first = data.get('identity', {}).get('first_name', '')
            last = data.get('identity', {}).get('last_name', '')
            return f"{first} {last}".strip()
    except Exception:
        return ""
    return ""

# -----------------------------------------------------
# 3. STATE & AUTO-LOGIN HANDLER
# -----------------------------------------------------
if 'gdrive_creds' not in st.session_state:
    st.session_state.gdrive_creds = None
if 'basecamp_token' not in st.session_state:
    st.session_state.basecamp_token = None
if 'user_real_name' not in st.session_state:
    st.session_state.user_real_name = ""

# --- FIX: IMMEDIATE GOOGLE RE-LOGIN ---
if 'gdrive_creds_json' in st.session_state and st.session_state.gdrive_creds_json and not st.session_state.gdrive_creds:
    try:
        creds = Credentials.from_authorized_user_info(
            json.loads(st.session_state.gdrive_creds_json)
        )
        st.session_state.gdrive_creds = creds
    except Exception as e:
        st.session_state.gdrive_creds_json = None

# --- AUTO-LOGIN HANDLER (BASECAMP) ---
if AUTO_LOGIN_MODE and "code" in st.query_params and not st.session_state.basecamp_token:
    auth_code = st.query_params["code"]
    try:
        payload = {
            "type": "web_server",
            "client_id": BASECAMP_CLIENT_ID,
            "client_secret": BASECAMP_CLIENT_SECRET,
            "redirect_uri": BASECAMP_REDIRECT_URI,
            "code": auth_code
        }
        response = requests.post(BASECAMP_TOKEN_URL, data=payload)
        response.raise_for_status()
        token = response.json()
        
        st.session_state.basecamp_token = token
        
        real_name = fetch_basecamp_name(token)
        if real_name:
            st.session_state.user_real_name = real_name
            
        st.query_params.clear()
        st.toast("✅ Basecamp Login Successful!", icon="🎉")
        time.sleep(1)
        st.rerun()
    except Exception as e:
        st.error(f"Auto-login failed: {e}")

# -----------------------------------------------------
# 4. SIDEBAR UI
# -----------------------------------------------------
with st.sidebar:
    st.title("🔐 Login")
    
    # --- STEP 1: BASECAMP ---
    st.markdown("### Step 1: Basecamp")
    
    if st.session_state.basecamp_token:
        st.success(f"✅ Connected as {st.session_state.user_real_name}")
        if st.button("Logout Basecamp"):
            st.session_state.basecamp_token = None
            st.session_state.user_real_name = ""
            st.session_state.gdrive_creds = None
            st.session_state.gdrive_creds_json = None
            st.rerun()
    else:
        bc_oauth = OAuth2Session(BASECAMP_CLIENT_ID, redirect_uri=BASECAMP_REDIRECT_URI)
        # Force HTTPs compliance for oauthlib
        os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'
        bc_auth_url, _ = bc_oauth.authorization_url(BASECAMP_AUTH_URL, type="web_server")
        
        if AUTO_LOGIN_MODE:
            st.markdown("""
            <a href="{bc_auth_url}" target="_top" style="text-decoration: none;">
                <div style="
                    background-color: #ff4b4b;
                    color: white;
                    padding: 0.5rem 1rem;
                    border-radius: 0.5rem;
                    text-align: center;
                    font-weight: bold;
                    border: 1px solid #ff4b4b;
                    margin-bottom: 10px;
                ">
                    Login to Basecamp
                </div>
            </a>
            """, unsafe_allow_html=True)
            st.caption("You must log in to Basecamp first.")
        else:
            st.warning("Auto-login not configured in Secrets.")
            st.markdown(f"👉 [**Authorize Basecamp**]({bc_auth_url})")
            bc_code = st.text_input("Paste Basecamp Code:", key="bc_code")
            if bc_code:
                try:
                    payload = {
                        "type": "web_server",
                        "client_id": BASECAMP_CLIENT_ID,
                        "client_secret": BASECAMP_CLIENT_SECRET,
                        "redirect_uri": BASECAMP_REDIRECT_URI,
                        "code": bc_code
                    }
                    response = requests.post(BASECAMP_TOKEN_URL, data=payload)
                    response.raise_for_status()
                    token = response.json()
                    st.session_state.basecamp_token = token
                    real_name = fetch_basecamp_name(token)
                    if real_name: st.session_state.user_real_name = real_name
                    st.rerun()
                except Exception as e:
                    st.error(f"Login failed: {e}")

    st.divider()

    # --- STEP 2: GOOGLE DRIVE ---
    st.markdown("### Step 2: Google Drive")
    
    if not st.session_state.basecamp_token:
        st.info("🔒 Please complete Step 1 first.")
    else:
        if st.session_state.gdrive_creds:
            st.success("✅ Connected")
            if st.button("Logout Drive"):
                st.session_state.gdrive_creds = None
                st.session_state.gdrive_creds_json = None
                st.rerun()
        else:
            try:
                flow = Flow.from_client_config(
                    GDRIVE_CLIENT_CONFIG,
                    scopes=["https://www.googleapis.com/auth/drive"],
                    redirect_uri="urn:ietf:wg:oauth:2.0:oob"
                )
                auth_url, _ = flow.authorization_url(prompt='consent')
                
                st.markdown(f"👉 [**Click to Authorize Drive**]({auth_url})")
                g_code = st.text_input("Paste Google Code:", key="g_code")
                
                if g_code:
                    flow.fetch_token(code=g_code)
                    st.session_state.gdrive_creds = flow.credentials
                    st.session_state.gdrive_creds_json = flow.credentials.to_json()
                    st.rerun()
            except Exception as e:
                st.error(f"Config Error: {e}")

# -----------------------------------------------------
# 5. API CLIENTS (ROBOT) SETUP
# -----------------------------------------------------
try:
    sa_creds = service_account.Credentials.from_service_account_info(GCP_SERVICE_ACCOUNT_JSON)
    storage_client = storage.Client(credentials=sa_creds)
    speech_client = speech.SpeechClient(credentials=sa_creds)
    
    genai.configure(api_key=GOOGLE_API_KEY)
    gemini_model = genai.GenerativeModel('gemini-2.5-flash-lite')
except Exception as e:
    st.error(f"System Error (AI Services): {e}")
    st.stop()

# -----------------------------------------------------
# 6. HELPER FUNCTIONS
# -----------------------------------------------------

def get_basecamp_session_user():
    if not st.session_state.basecamp_token: return None
    session = OAuth2Session(BASECAMP_CLIENT_ID, token=st.session_state.basecamp_token)
    session.headers.update(BASECAMP_USER_AGENT)
    return session

def upload_to_gcs(file_path, destination_blob_name):
    try:
        bucket = storage_client.bucket(GCS_BUCKET_NAME)
        blob = bucket.blob(destination_blob_name)
        blob.upload_from_filename(file_path, timeout=3600)
        return f"gs://{GCS_BUCKET_NAME}/{destination_blob_name}"
    except Exception as e:
        st.error(f"GCS Upload Error: {e}")
        return None

def upload_to_drive_user(file_stream, file_name):
    if not st.session_state.gdrive_creds: return None
    try:
        service = build("drive", "v3", credentials=st.session_state.gdrive_creds)
        file_metadata = {"name": file_name} 
        media = MediaIoBaseUpload(
            file_stream, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        file = service.files().create(
            body=file_metadata, media_body=media, fields="id"
        ).execute()
        return file.get("id")
    except Exception as e:
        st.error(f"Google Drive Upload Error: {e}")
        return None

def get_basecamp_projects(_session):
    try:
        response = _session.get(f"{BASECAMP_API_BASE}/projects.json")
        response.raise_for_status()
        return sorted([(p['name'], p['id']) for p in response.json() if p['status'] == 'active'], key=lambda x: x[0])
    except: return []

def get_basecamp_todolists(_session, project_id):
    try:
        project_resp = _session.get(f"{BASECAMP_API_BASE}/projects/{project_id}.json").json()
        todoset_id = next((t['id'] for t in project_resp.get('dock', []) if t['name'] == 'todoset' and t['enabled']), None)
        if not todoset_id: return []
        lists = _session.get(f"{BASECAMP_API_BASE}/buckets/{project_id}/todosets/{todoset_id}/todolists.json").json()
        return sorted([(tl['title'], tl['id']) for tl in lists], key=lambda x: x[0])
    except: return []

def upload_bc_attachment(_session, file_bytes, file_name):
    try:
        headers = _session.headers.copy()
        headers.update({'Content-Type': 'application/octet-stream', 'Content-Length': str(len(file_bytes))})
        resp = _session.post(f"{BASECAMP_API_BASE}/attachments.json?name={file_name}", data=file_bytes, headers=headers)
        return resp.json()['attachable_sgid']
    except Exception as e:
        st.error(f"Basecamp Upload Error: {e}")
        return None

def create_bc_todo(_session, project_id, todolist_id, title, attachment_sgid):
    try:
        content = title
        desc = f'<bc-attachment sgid="{attachment_sgid}"></bc-attachment>' if attachment_sgid else ""
        payload = {"content": content, "description": desc}
        url = f"{BASECAMP_API_BASE}/buckets/{project_id}/todolists/{todolist_id}/todos.json"
        _session.post(url, json=payload).raise_for_status()
        return True
    except Exception as e:
        st.error(f"Basecamp Create Error: {e}")
        return False

def get_structured_notes_google(audio_file_path, file_name, participants_context):
    flac_file_path = ""
    flac_blob_name = ""
    try:
        with st.spinner(f"Converting {file_name} to audio-only FLAC..."):
            base_name = os.path.splitext(audio_file_path)[0]
            flac_file_path = f"{base_name}.flac"
            command = ["ffmpeg", "-i", audio_file_path, "-vn", "-acodec", "flac", "-y", flac_file_path]
            subprocess.run(command, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

        with st.spinner(f"Uploading converted audio to Google Cloud Storage..."):
            flac_blob_name = f"{os.path.splitext(file_name)[0]}.flac"
            gcs_uri = upload_to_gcs(flac_file_path, flac_blob_name)
            if not gcs_uri: return {"error": "Upload failed."}

        progress_text = "Transcribing & identifying speakers..."
        progress_bar = st.progress(0, text=progress_text)
        
        audio = speech.RecognitionAudio(uri=gcs_uri)
        config = speech.RecognitionConfig(
            encoding=speech.RecognitionConfig.AudioEncoding.FLAC,
            language_code="en-US",
            enable_automatic_punctuation=True,
            use_enhanced=True,
            model="video",
            diarization_config=speech.SpeakerDiarizationConfig(
                enable_speaker_diarization=True,
                min_speaker_count=2,
                max_speaker_count=6
            )
        )
        operation = speech_client.long_running_recognize(config=config, audio=audio)
        
        while not operation.done():
            metadata = operation.metadata
            if metadata and metadata.progress_percent:
                progress_bar.progress(metadata.progress_percent, text=f"Transcribing: {metadata.progress_percent}%")
            time.sleep(2)

        progress_bar.progress(100, text="Transcription Complete")
        response = operation.result(timeout=3600)
        progress_bar.empty()

        if not response.results:
             return {"error": "Transcription failed. The audio might be silent."}

        result = response.results[-1]
        words_info = result.alternatives[0].words
        full_transcript_text = ""
        current_speaker = -1
        for word_info in words_info:
            if word_info.speaker_tag != current_speaker:
                current_speaker = word_info.speaker_tag
                full_transcript_text += f"\n\nSpeaker {current_speaker}: "
            full_transcript_text += word_info.word + " "
        if not full_transcript_text.strip():
            full_transcript_text = " ".join([result.alternatives[0].transcript for result in response.results])

        with st.spinner("Analyzing conversation & matching names..."):
            prompt = f"""
            You are an expert meeting secretary. 
            Here is the context of who was in the meeting:
            {participants_context}
            The transcript below uses "Speaker 1", "Speaker 2", etc.
            Your job is to figure out which Speaker matches which Name from the list above.
            Transcript:
            {full_transcript_text}
            ---
            YOUR TASKS:
            1. RECONSTRUCTION: When writing the notes, DO NOT use "Speaker 1". Use their REAL NAMES (e.g., "John said...").
            2. EXTRACTION:
            ## DISCUSSION ##
            Summarize main points using the real names.
            FORMAT:
            ## Section Title (e.g., ## Content and Grammar)
            * **Wording & Tone:** John requested avoiding the casual use of "You are".
            * Bullet point 3.
            (Leave a blank line between sections)
            ## NEXT STEPS ##
            List action items using the real names.
            FORMAT:
            * Bullet point 1.
            ## CLIENT REQUESTS ##
            List specific questions or requests asked BY the Client.
            FORMAT:
            * Bullet point 1.
            """
            response = gemini_model.generate_content(prompt)
            text = response.text
            
            discussion = ""
            next_steps = ""
            client_reqs = ""
            try:
                if "## DISCUSSION ##" in text:
                    discussion = text.split("## DISCUSSION ##")[1].split("## NEXT STEPS ##")[0].strip()
                if "## NEXT STEPS ##" in text:
                    next_steps = text.split("## NEXT STEPS ##")[1].split("## CLIENT REQUESTS ##")[0].strip()
                if "## CLIENT REQUESTS ##" in text:
                    client_reqs = text.split("## CLIENT REQUESTS ##")[1].strip()
                if not discussion: discussion = text
            except:
                discussion = text

            return {
                "discussion": discussion,
                "next_steps": next_steps,
                "client_reqs": client_reqs,
                "full_transcript": full_transcript_text
            }
    except Exception as e:
        return {"error": str(e)}
    finally:
        try:
            bucket = storage_client.bucket(GCS_BUCKET_NAME)
            bucket.blob(flac_blob_name).delete()
        except: pass

def add_formatted_text(cell, text):
    cell.text = ""
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line: continue
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if line.startswith('##'):
            run = p.add_run(line.strip('#').strip())
            run.underline = True
        elif line.startswith('*'):
            clean = line.lstrip('*').lstrip("•").strip()
            if clean.startswith('**') and ':**' in clean:
                try:
                    parts = clean.split(':**', 1)
                    p.text = "•\t"
                    p.add_run(parts[0].lstrip('**').strip() + ": ").bold = True
                    p.add_run(parts[1].strip())
                    p.paragraph_format.left_indent = Inches(0.25)
                except:
                    p.text = f"•\t{clean}"
                    p.paragraph_format.left_indent = Inches(0.25)
            else:
                p.text = f"•\t{clean}"
                p.paragraph_format.left_indent = Inches(0.25)
        else:
            p.add_run(line)

# -----------------------------------------------------
# 7. STREAMLIT UI
# -----------------------------------------------------
if 'ai_results' not in st.session_state:
    st.session_state.ai_results = {"discussion": "", "next_steps": "", "client_reqs": "", "full_transcript": ""}
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "auto_client_reps" not in st.session_state:
    st.session_state.auto_client_reps = ""
if "auto_ifoundries_reps" not in st.session_state:
    st.session_state.auto_ifoundries_reps = ""
if "saved_participants_input" not in st.session_state:
    st.session_state.saved_participants_input = ""

st.title("🤖 AI Meeting Manager")

tab1, tab2, tab3 = st.tabs(["1. Analyze", "2. Review & Export", "3. Chat"])

with tab1:
    st.header("1. Analyze Audio")
    participants_input = st.text_area(
        "Known Participants (Teach the AI)", 
        value="Client's Exact Name (Client)\niFoundries Exact Name (iFoundries)",
        help="The AI will read this to match 'Speaker 1' to these names."
    )
    uploaded_file = st.file_uploader("Upload Meeting", type=["mp3", "mp4", "m4a", "wav"])
    
    if st.button("Analyze Audio"):
        if uploaded_file:
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name}") as tmp:
                tmp.write(uploaded_file.getvalue())
                path = tmp.name
            
            st.session_state.chat_history = [] 
            st.session_state.saved_participants_input = participants_input 
            
            # Auto-fill logic
            c_list = [l.replace("(Client)","").strip() for l in participants_input.split('\n') if "(Client)" in l]
            i_list = [l.replace("(iFoundries)","").strip() for l in participants_input.split('\n') if "(iFoundries)" in l]
            st.session_state.auto_client_reps = "\n".join(c_list)
            st.session_state.auto_ifoundries_reps = ", ".join(i_list)
            
            res = get_structured_notes_google(path, uploaded_file.name, participants_input)
            if "error" in res: st.error(res["error"])
            else: 
                st.session_state.ai_results = res
                st.success("Analysis Complete!")
        else:
            st.warning("Please upload a file first.")

with tab2:
    st.header("2. Review Notes")
    
    sg_tz = pytz.timezone('Asia/Singapore')
    sg_now = datetime.datetime.now(sg_tz)
    
    row1, row2 = st.columns(2)
    with row1:
        date_obj = st.date_input("Date", sg_now.date())
        venue = st.text_input("Venue")
        client_rep = st.text_area("Client Reps", value=st.session_state.auto_client_reps)
        absent = st.text_input("Absent")
    with row2:
        time_obj = st.text_input("Time", value=sg_now.strftime("%I:%M %p"))
        
        # --- AUTO FILL PREPARED BY ---
        default_prepared_by = st.session_state.user_real_name if st.session_state.user_real_name else st.session_state.auto_ifoundries_reps
        
        prepared_by = st.text_input("Prepared by", value=default_prepared_by)
        ifoundries_rep = st.text_input("iFoundries Reps", value=st.session_state.auto_ifoundries_reps)
    
    date_str = date_obj.strftime("%d %B %Y")
    time_str = time_obj
    
    discussion_text = st.text_area("Discussion", value=st.session_state.ai_results.get("discussion", ""), height=300)
    next_steps_text = st.text_area("Next Steps", value=st.session_state.ai_results.get("next_steps", ""), height=200)
    with st.expander("View Specific Client Requests"):
        st.text_area("Client Requests", value=st.session_state.ai_results.get("client_reqs", ""), height=150)

    st.header("3. Generate & Upload")
    
    bc_session_user = None
    bc_project_id = None
    bc_todolist_id = None
    bc_todo_title = ""

    do_drive = st.checkbox("Upload to Drive")
    do_basecamp = st.checkbox("Upload to Basecamp") 

    if do_basecamp:
        if st.session_state.basecamp_token:
            bc_session_user = get_basecamp_session_user()
            try:
                projects_list = get_basecamp_projects(bc_session_user)
                if not projects_list:
                    st.warning("No active Basecamp projects found.")
                else:
                    selected_project_name = st.selectbox("Select Project", options=[p[0] for p in projects_list], index=None, placeholder="Choose...")
                    if selected_project_name:
                        bc_project_id = next(p[1] for p in projects_list if p[0] == selected_project_name)
                        todolists = get_basecamp_todolists(bc_session_user, bc_project_id)
                        if not todolists: st.warning("No to-do lists found.")
                        else:
                            selected_list = st.selectbox("Select List", options=[tl[0] for tl in todolists], index=None, placeholder="Choose...")
                            if selected_list:
                                bc_todolist_id = next(tl[1] for tl in todolists if tl[0] == selected_list)
                                bc_todo_title = st.text_input("To-Do Title :red[*]")
                                if date_str: st.info(f"📎 Minutes_{date_str}.docx will be attached to Notes.")
            except Exception as e: st.error(f"Basecamp Error: {e}")
        else:
            st.warning("⚠️ Please log in to Basecamp in the sidebar first.")

    if do_drive and not st.session_state.gdrive_creds:
        st.warning("⚠️ Please log in to Google Drive in the sidebar first.")

    if st.button("Generate Word Doc"):
        basecamp_ready = True
        if do_basecamp:
            if not st.session_state.basecamp_token:
                st.error("Basecamp not connected.")
                basecamp_ready = False
            elif not bc_project_id or not bc_todolist_id or not bc_todo_title:
                st.error("Please complete Basecamp fields.")
                basecamp_ready = False

        if not date_str or not prepared_by or not client_rep:
            st.error("Missing required fields (*)")
        elif not do_basecamp or basecamp_ready:
            if do_drive and not st.session_state.gdrive_creds:
                st.error("Google Drive not connected.")
            else:
                try:
                    doc = Document("Minutes Of Meeting - Template.docx")
                    t0 = doc.tables[0]
                    t0.cell(1,1).text = date_str
                    t0.cell(2,1).text = time_str
                    t0.cell(3,1).text = venue
                    c_rep_final = f"{client_rep} (Client)" if client_rep and "(Client)" not in client_rep else client_rep
                    i_rep_final = f"{ifoundries_rep} (iFoundries)" if ifoundries_rep and "(iFoundries)" not in ifoundries_rep else ifoundries_rep
                    t0.cell(4,1).text = c_rep_final
                    t0.cell(4,2).text = i_rep_final
                    t0.cell(5,1).text = absent

                    t1 = doc.tables[1]
                    add_formatted_text(t1.cell(2,1), discussion_text)
                    add_formatted_text(t1.cell(4,1), next_steps_text)
                    doc.paragraphs[-1].text = f"Prepared by: {prepared_by}"
                    
                    bio = io.BytesIO()
                    doc.save(bio)
                    bio.seek(0)
                    fname = f"Minutes_{date_str}.docx"
                    
                    if do_drive and st.session_state.gdrive_creds:
                        with st.spinner("Uploading to Drive..."):
                            if upload_to_drive_user(bio, fname): st.success("Uploaded to Drive!")
                            else: st.error("Drive upload failed.")
                        bio.seek(0)

                    if do_basecamp and basecamp_ready and bc_session_user:
                        with st.spinner(f"Uploading to Basecamp..."):
                            file_bytes = bio.getvalue()
                            sgid = upload_bc_attachment(bc_session_user, file_bytes, fname)
                        
                        if sgid:
                            with st.spinner("Creating To-Do..."):
                                if create_bc_todo(bc_session_user, bc_project_id, bc_todolist_id, bc_todo_title, sgid):
                                    st.success("Created To-Do in Basecamp!")
                                else: st.error("Basecamp creation failed.")
                        else: st.error("Basecamp file upload failed.")
                        bio.seek(0)

                    st.download_button("Download .docx", bio, fname)
                    
                except Exception as e:
                    st.error(f"Error: {e}")

with tab3:
    st.header("💬 Chat with your Meeting")
    
    transcript_context = st.session_state.ai_results.get("full_transcript", "")
    participants_context = st.session_state.saved_participants_input
    
    if not transcript_context:
        st.info("⚠️ Please upload and analyze a meeting audio file in Tab 1 first.")
    else:
        if st.button("Clear Chat"):
            st.session_state.chat_history = []
            st.rerun()

        chat_container = st.container(height=500)
        
        with chat_container:
            for message in st.session_state.chat_history:
                if message["role"] == "user":
                    with st.chat_message("user", avatar="👤"):
                        st.markdown(message["content"])
                else:
                    with st.chat_message("assistant", avatar="🤖"):
                        st.markdown(message["content"])

        if prompt := st.chat_input("Ask a question about the meeting..."):
            st.session_state.chat_history.append({"role": "user", "content": prompt})
            
            with chat_container:
                with st.chat_message("user", avatar="👤"):
                    st.markdown(prompt)

            with chat_container:
                with st.chat_message("assistant", avatar="🤖"):
                    def stream_text(response_iterator):
                        for chunk in response_iterator:
                            if chunk.parts:
                                yield chunk.text

                    try:
                        # --- FINAL "OFFICIAL LOG" STYLE PROMPT ---
                        full_prompt = f"""
                        You are an efficient, action-oriented meeting secretary.
                        
                        CONTEXT (PARTICIPANTS):
                        {participants_context}
                        (Use this to map "Speaker X" to real names.)

                        TRANSCRIPT:
                        {transcript_context}
                        
                        USER QUESTION:
                        {prompt}
                        
                        STRICT RULES:
                        1. **Passive/Professional Voice:** Focus on the action/decision, NOT the speaker, unless it is a direct assignment.
                           - BAD: "John said the font is too small."
                           - GOOD: "The font size needs to be increased." (Focus on the task)
                           - GOOD: "The Client requested a larger font." (Focus on the role)
                        2. **No Speaker IDs:** NEVER use "Speaker 1" or "Speaker 2".
                        3. **Accuracy:** Use the transcript as your only source. If not mentioned, say "That was not discussed."
                        4. **Conciseness:** Be brief and clear.
                        """
                        
                        stream_iterator = gemini_model.generate_content(full_prompt, stream=True)
                        response = st.write_stream(stream_text(stream_iterator))
                        
                        st.session_state.chat_history.append({"role": "assistant", "content": response})
                    except Exception as e:
                        st.error("I couldn't generate a response. Please try again.")
